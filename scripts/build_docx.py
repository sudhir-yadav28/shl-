"""Convert APPROACH.md to APPROACH.docx with sensible defaults.

Targets a 2-page printable document. Uses Calibri 11pt body, 14pt headings,
narrow margins, and inline-code styling for backticks. Not a full Markdown
renderer — just the constructs APPROACH.md actually uses (headings,
paragraphs, bold, inline code, horizontal rule, link auto-detection).

Usage:
    source .venv/bin/activate
    python scripts/build_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches


SRC = Path("APPROACH.md")
OUT = Path("APPROACH.docx")


def add_inline_runs(paragraph, text: str) -> None:
    """Tokenize a line of markdown into runs.

    Recognized: **bold**, *italic*, `code`, [link text](url).
    Everything else becomes a plain run.
    """
    # Match in priority order; whichever pattern starts earliest wins.
    pattern = re.compile(
        r"(\*\*([^*]+)\*\*)"           # group 1,2 — bold
        r"|(`([^`]+)`)"                # group 3,4 — code
        r"|(\[([^\]]+)\]\(([^)]+)\))"  # group 5,6,7 — link
        r"|(\*([^*]+)\*)"              # group 8,9 — italic
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(4))
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        elif m.group(5):
            link_text, url = m.group(6), m.group(7)
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0x03, 0x66, 0xD6)
            run.underline = True
        elif m.group(8):
            run = paragraph.add_run(m.group(9))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main() -> None:
    md = SRC.read_text(encoding="utf-8")

    doc = Document()

    # Narrow margins so the prose fits in ~2 pages
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Horizontal rule -> a small spacer paragraph
        if re.match(r"^-{3,}$", line):
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            i += 1
            continue

        # Default — paragraph (collapse soft-wrapped lines into one)
        para_lines = [line]
        i += 1
        while (
            i < len(lines)
            and lines[i].strip()
            and not lines[i].startswith("#")
            and not lines[i].startswith("- ")
            and not lines[i].startswith("* ")
            and not re.match(r"^-{3,}$", lines[i])
        ):
            para_lines.append(lines[i].rstrip())
            i += 1
        merged = " ".join(para_lines).strip()
        if merged:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            add_inline_runs(p, merged)

    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"wrote {OUT}  ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
